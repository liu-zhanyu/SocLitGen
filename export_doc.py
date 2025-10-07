from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
import markdown
from bs4 import BeautifulSoup
from pymongo import MongoClient
from typing import Optional, Dict, List
import re


class ReviewDocumentExporter:
    """文献综述Word文档导出器 - 从literature_data直接生成APA格式参考文献"""

    def __init__(self, mongo_uri: str):
        """
        初始化导出器

        Args:
            mongo_uri: MongoDB连接字符串
        """
        self.mongo_client = MongoClient(mongo_uri)
        self.db = self.mongo_client["Newbit"]
        self.collection = self.db["review"]

    def export_review_document(
            self,
            user_id: str,
            task_id: str,
            output_dir: str
    ) -> Optional[str]:
        """
        导出文献综述为Word文档

        Args:
            user_id: 用户ID
            task_id: 任务ID
            output_dir: 输出目录路径

        Returns:
            str: 导出的文件完整路径，失败返回None
        """
        # 从MongoDB获取数据
        review_doc = self.collection.find_one({
            "user_id": user_id,
            "task_id": task_id
        })

        if not review_doc:
            print(f"未找到文档: user_id={user_id}, task_id={task_id}")
            return None

        review_text = review_doc.get("review_text", {})
        structure = review_doc.get("structure", 0)

        # 获取主题作为文件名
        main_topic = review_text.get("main_topic", "文献综述")

        # 清理文件名中的非法字符
        safe_filename = self._sanitize_filename(main_topic)
        output_path = f"{output_dir}/{safe_filename}.docx"

        # 创建Word文档
        doc = Document()
        self._set_document_style(doc)

        # 添加概述
        overview = review_text.get("overview", "")
        if overview:
            self._add_markdown_content(doc, overview)

        # 根据结构类型添加主体内容
        topics = review_text.get("topics", {})
        if structure == 1:
            # 中英文分离结构
            self._add_structure_1_content(doc, topics)
        else:
            # 标准结构
            self._add_structure_0_content(doc, topics)

        # 添加研究评述
        critique = review_text.get("critique", "")
        if critique:
            heading = doc.add_heading("研究评述", level=2)
            self._apply_heading_style(heading, level=2)
            self._add_markdown_content(doc, critique)

        # 添加参考文献（从literature_data生成）
        literature_data = review_doc.get("literature_data", {})
        if literature_data:
            self._add_references_from_literature(doc, literature_data, topics)

        # 保存文档
        doc.save(output_path)
        print(f"文档已保存至: {output_path}")
        return output_path

    def _add_references_from_literature(
            self,
            doc: Document,
            literature_data: Dict,
            topics_text: Dict
    ):
        """
        从literature_data生成参考文献列表（统一APA格式）

        Args:
            doc: Word文档对象
            literature_data: 文献数据
            topics_text: 主题文本（用于检查引用）
        """
        # 收集所有被引用的文献
        cited_papers = []
        combined_text = ""

        # 合并所有主题文本
        if isinstance(topics_text, dict):
            if "topics_zh" in topics_text and "topics_en" in topics_text:
                # Structure 1
                combined_text = "\n".join(topics_text["topics_zh"].values())
                combined_text += "\n" + "\n".join(topics_text["topics_en"].values())
            else:
                # Structure 0
                combined_text = "\n".join(topics_text.values())

        # 遍历文献数据，找出被引用的文献
        for topic, content_types in literature_data.items():
            for content_type, papers in content_types.items():
                for paper in papers:
                    # 检查是否被引用
                    if self._is_paper_cited(paper, combined_text):
                        cited_papers.append(paper)

        if not cited_papers:
            return

        # 去重（基于authors和year）
        unique_papers = self._deduplicate_papers(cited_papers)

        # 添加参考文献标题
        ref_heading = doc.add_heading('参考文献', level=2)
        self._apply_heading_style(ref_heading, level=2)

        # 所有文献使用APA格式，按作者姓氏字母排序
        unique_papers.sort(key=lambda x: x.get('authors', '').lower())

        for paper in unique_papers:
            ref_text = self._format_apa_reference(paper)
            self._add_reference_paragraph(doc, ref_text)

    def _is_paper_cited(self, paper: Dict, text: str) -> bool:
        """检查文献是否在文本中被引用"""
        authors = paper.get('authors', '')
        year = int(paper.get('year', ''))

        if not authors or not year:
            return False

        # 提取第一作者姓氏
        first_author = authors.split(';')[0].strip()
        if ',' in first_author:
            surname = first_author.split(',')[0].strip()
        else:
            surname = first_author
        # 检查是否出现"作者姓氏"和"年份"
        return surname in text and str(year) in text

    def _deduplicate_papers(self, papers: List[Dict]) -> List[Dict]:
        """去重文献列表"""
        seen = set()
        unique = []

        for paper in papers:
            key = (paper.get('authors', ''), paper.get('year', ''))
            if key not in seen:
                seen.add(key)
                unique.append(paper)

        return unique

    def _format_apa_reference(self, paper: Dict) -> str:
        """
        格式化为APA格式参考文献
        """
        parts = []

        # 1. 作者 - 增强过滤
        authors = paper.get('authors', '')
        if authors and authors.strip():
            # 分割并严格过滤
            author_list = [a.strip() for a in authors.split(';')]
            author_list = [a for a in author_list if a and a.strip() and not all(c in '.,; ' for c in a)]

            if author_list:
                formatted_authors = self._format_apa_authors(author_list)
                parts.append(formatted_authors)
            else:
                parts.append("[作者信息缺失]")
        else:
            parts.append("[作者信息缺失]")

        # 2. 年份
        year = paper.get('year', '')
        if year:
            parts.append(f"({year}).")

        # 3. 标题（sentence case）
        title = paper.get('title', '')
        if title:
            title = self._to_sentence_case(title)
            parts.append(f"{title}.")

        # 4. 期刊名（title case）+ 卷期页
        journal = paper.get('journal', '')
        if journal:
            journal = self._to_title_case(journal)
            journal_part = journal

            volume = paper.get('vO', '') or paper.get('volume', '')
            issue = paper.get('issue', '')
            pages = paper.get('page_range', '')

            if volume:
                journal_part += f", {volume}"
                if issue:
                    journal_part += f"({issue})"

            if pages:
                journal_part += f", {pages}"

            parts.append(f"{journal_part}.")

        return " ".join(parts)

    def _format_apa_authors(self, author_list: List[str]) -> str:
        """格式化APA作者名单"""
        if not author_list:
            return "[作者信息缺失]"

        formatted = []
        for author in author_list:
            # 跳过空作者名
            if not author or not author.strip():
                continue

            author = author.strip()

            # 如果是纯标点符号或空字符，跳过
            if not author or all(c in '.,; ' for c in author):
                continue

            if ',' in author:
                # 英文作者: Last, F. M.
                parts = author.split(',')
                last_name = parts[0].strip()
                first_names = parts[1].strip() if len(parts) > 1 else ""

                # 跳过空的姓氏
                if not last_name:
                    continue

                # 提取首字母
                initials = []
                for name in first_names.split():
                    if name and name.strip():
                        initials.append(name[0].upper() + '.')

                if initials:
                    formatted_author = f"{last_name}, {' '.join(initials)}"
                else:
                    formatted_author = last_name
            else:
                # 中文作者或其他格式 - 直接使用
                formatted_author = author

            # 再次检查确保作者名不为空且有意义
            if formatted_author and formatted_author.strip():
                formatted.append(formatted_author.strip())

        # 如果所有作者都被过滤掉了
        if not formatted:
            return "[作者信息缺失]"

        # 连接作者
        if len(formatted) == 1:
            author_str = formatted[0]
            # 确保以句点结尾
            return author_str if author_str.endswith('.') else author_str + '.'
        elif len(formatted) == 2:
            author1 = formatted[0]
            author2 = formatted[1]
            # 确保第二个作者以句点结尾
            if not author2.endswith('.'):
                author2 += '.'
            return f"{author1}, & {author2}"
        else:
            # 多个作者
            author_last = formatted[-1]
            other_authors = formatted[:-1]

            # 确保最后一个作者以句点结尾
            if not author_last.endswith('.'):
                author_last += '.'

            return ", ".join(other_authors) + f", & {author_last}"

    def _to_sentence_case(self, text: str) -> str:
        """
        转换为句子式大小写（sentence case）
        只有首字母、专有名词、缩写词大写
        """
        if not text:
            return text

        words = text.split()
        result = []

        for i, word in enumerate(words):
            # 保留全大写的缩写词（2-5个字母）
            if word.isupper() and 2 <= len(word) <= 5:
                result.append(word)
            # 保留带连字符的专有名词
            elif '-' in word and any(p[0].isupper() for p in word.split('-') if p):
                result.append(word)
            # 第一个词首字母大写
            elif i == 0:
                result.append(word[0].upper() + word[1:].lower() if len(word) > 1 else word.upper())
            # 冒号后的词首字母大写
            elif result and result[-1].endswith(':'):
                result.append(word[0].upper() + word[1:].lower() if len(word) > 1 else word.upper())
            # 其他词小写
            else:
                result.append(word.lower())

        return ' '.join(result)

    def _to_title_case(self, text: str) -> str:
        """
        转换为标题式大小写（title case）
        主要单词首字母大写，小词（冠词、介词等）小写
        """
        if not text:
            return text

        # 小词列表
        small_words = {
            'a', 'an', 'and', 'as', 'at', 'but', 'by', 'for',
            'in', 'of', 'on', 'or', 'the', 'to', 'up', 'with'
        }

        words = text.split()
        result = []

        for i, word in enumerate(words):
            # 第一个和最后一个词总是首字母大写
            if i == 0 or i == len(words) - 1:
                result.append(word[0].upper() + word[1:].lower() if len(word) > 1 else word.upper())
            # 小词保持小写
            elif word.lower() in small_words:
                result.append(word.lower())
            # 其他词首字母大写
            else:
                result.append(word[0].upper() + word[1:].lower() if len(word) > 1 else word.upper())

        return ' '.join(result)

    def _add_reference_paragraph(self, doc: Document, text: str):
        """添加一条参考文献（APA格式，无序号）"""
        paragraph = doc.add_paragraph()
        paragraph.clear()
        run = paragraph.add_run(text)

        # 参考文献使用悬挂缩进格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(24)
        paragraph_format.first_line_indent = Pt(-24)
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # 设置字体
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，移除非法字符"""
        illegal_chars = r'[<>:"/\\|?*]'
        safe_name = re.sub(illegal_chars, '_', filename)
        safe_name = safe_name.strip('. ')

        if len(safe_name) > 100:
            safe_name = safe_name[:100]

        if not safe_name:
            safe_name = "文献综述"

        return safe_name

    def _add_structure_0_content(self, doc: Document, topics: Dict):
        """添加标准结构内容"""
        if isinstance(topics, dict):
            for topic_name, topic_content in topics.items():
                if isinstance(topic_content, str):
                    self._add_markdown_content(doc, topic_content)

    def _add_structure_1_content(self, doc: Document, topics: Dict):
        """添加中英文分离结构内容"""
        topics_zh = topics.get("topics_zh", {})
        if topics_zh:
            heading = doc.add_heading("一、国内相关研究", level=1)
            self._apply_heading_style(heading, level=1)

            for topic_name, topic_content in topics_zh.items():
                if isinstance(topic_content, str):
                    self._add_markdown_content(doc, topic_content)

        topics_en = topics.get("topics_en", {})
        if topics_en:
            heading = doc.add_heading("二、国外相关研究", level=1)
            self._apply_heading_style(heading, level=1)

            for topic_name, topic_content in topics_en.items():
                if isinstance(topic_content, str):
                    self._add_markdown_content(doc, topic_content)

    def _add_markdown_content(self, doc: Document, markdown_text: str):
        """
        解析并添加Markdown内容到Word文档

        Args:
            doc: Word文档对象
            markdown_text: Markdown格式文本
        """
        if not markdown_text:
            return

        # 按行处理
        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # 检测表格
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                # 收集完整表格
                table_lines = [line]
                j = i + 1
                while j < len(lines) and '|' in lines[j]:
                    table_lines.append(lines[j])
                    j += 1

                # 提取表格标题（如果有）
                table_title = ""
                if i > 0 and lines[i - 1].strip():
                    prev_line = lines[i - 1].strip()
                    if not prev_line.startswith('#'):
                        table_title = prev_line

                # 添加表格
                table_markdown = '\n'.join(table_lines)
                self._add_markdown_table_to_doc(doc, table_markdown, table_title)
                i = j
                continue

            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()

                if title_text:
                    heading = doc.add_heading(title_text, level=min(level, 3))
                    self._apply_heading_style(heading, level=min(level, 3))
                i += 1
                continue

            # 处理普通段落
            if line.strip():
                paragraph = doc.add_paragraph(line.strip())
                self._apply_paragraph_style(paragraph)

            i += 1

    def _add_markdown_table_to_doc(
            self,
            doc: Document,
            markdown_table: str,
            table_title: str = ""
    ):
        """将Markdown表格添加到Word文档"""
        try:
            # 解析Markdown表格
            html = markdown.markdown(markdown_table, extensions=['tables'])
            soup = BeautifulSoup(html, 'html.parser')
            table_element = soup.find('table')

            if not table_element:
                return

            rows = table_element.find_all('tr')
            if not rows:
                return

            # 添加表格标题
            if table_title:
                title_paragraph = doc.add_paragraph()
                title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title_run = title_paragraph.add_run("表 " + table_title)
                title_font = title_run.font
                title_font.name = 'Times New Roman'
                title_font.size = Pt(10)
                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                title_format = title_paragraph.paragraph_format
                title_format.first_line_indent = Pt(0)
                title_format.line_spacing = 1

            # 创建Word表格
            num_cols = len(rows[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 填充表格内容
            for row_idx, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for col_idx, cell in enumerate(cells):
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        cell_text = cell.get_text(strip=True)
                        table.rows[row_idx].cells[col_idx].text = cell_text

                        # 设置单元格样式
                        cell_paragraph = table.rows[row_idx].cells[col_idx].paragraphs[0]
                        cell_paragraph.clear()
                        run = cell_paragraph.add_run(cell_text)

                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                        paragraph_format = cell_paragraph.paragraph_format
                        paragraph_format.first_line_indent = Pt(0)
                        paragraph_format.line_spacing = 1.0
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(0)

            # 设置表格边框
            self._set_table_borders(table)

        except Exception as e:
            print(f"添加表格失败: {e}")
            if table_title:
                paragraph = doc.add_paragraph(f"{table_title}\n{markdown_table}")
            else:
                paragraph = doc.add_paragraph(markdown_table)
            self._apply_paragraph_style(paragraph)

    def _set_document_style(self, doc: Document):
        """设置文档样式"""
        style = doc.styles['Normal']

        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        paragraph_format = style.paragraph_format
        paragraph_format.first_line_indent = Pt(24)
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _apply_paragraph_style(self, paragraph):
        """应用段落样式"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(24)
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for run in paragraph.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _apply_heading_style(self, heading, level: int = 1):
        """应用标题样式"""
        for run in heading.runs:
            font = run.font
            font.name = 'Times New Roman'
            if level == 1:
                font.size = Pt(16)
            elif level == 2:
                font.size = Pt(14)
            else:
                font.size = Pt(13)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _set_table_borders(self, table):
        """设置表格边框样式"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)

            tblBorders = OxmlElement('w:tblBorders')

            # 上边框
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), '18')
            top_border.set(qn('w:color'), '000000')
            tblBorders.append(top_border)

            # 下边框
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '18')
            bottom_border.set(qn('w:color'), '000000')
            tblBorders.append(bottom_border)

            # 其他边框设为无
            for border_name in ['left', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tblBorders.append(border)

            tblPr.append(tblBorders)

            # 设置第一行下边框
            if len(table.rows) > 0:
                self._set_row_bottom_border(table.rows[0], '9')

        except Exception as e:
            print(f"设置表格边框失败: {e}")

    def _set_row_bottom_border(self, row, border_size: str):
        """为指定行设置下边框"""
        try:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                existing_borders = tcPr.find(qn('w:tcBorders'))
                if existing_borders is not None:
                    tcPr.remove(existing_borders)

                tcBorders = OxmlElement('w:tcBorders')

                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'single')
                bottom_border.set(qn('w:sz'), border_size)
                bottom_border.set(qn('w:color'), '000000')
                tcBorders.append(bottom_border)

                for border_name in ['top', 'left', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)

                tcPr.append(tcBorders)

        except Exception as e:
            print(f"设置行边框失败: {e}")

    def close(self):
        """关闭数据库连接"""
        if self.mongo_client:
            self.mongo_client.close()


# 使用示例
if __name__ == "__main__":
    from components.config import *

    # 初始化导出器
    mongo_uri = f"mongodb://{MONGO_USER}:{MONGO_PASSWORD}@{MONGO_HOST}:{MONGO_PORT}"
    exporter = ReviewDocumentExporter(mongo_uri)

    # 导出文档
    success = exporter.export_review_document(
        user_id="string",
        task_id="review_7bda6029-f43e-42b2-9dfd-5a6f119a9d1f",
        output_dir=r"C:\Users\liuzh\Desktop"
    )

    if success:
        print(f"文档导出成功: {success}")
    else:
        print("文档导出失败!")

    # 关闭连接
    exporter.close()